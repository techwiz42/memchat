"""Conversations CRUD — list, create, delete, search, export user conversations."""

import io
import uuid
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy import select, delete, func, text
from sqlalchemy.ext.asyncio import AsyncSession

from auth.jwt import get_current_user_id
from models import Conversation, Message, get_db

router = APIRouter(prefix="/api/conversations", tags=["conversations"])


class ConversationOut(BaseModel):
    id: str
    title: str
    summary: str | None = None
    created_at: str
    updated_at: str


@router.get("", response_model=list[ConversationOut])
async def list_conversations(
    user_id: uuid.UUID = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """List user's conversations ordered by most recently updated."""
    result = await db.execute(
        select(Conversation)
        .where(Conversation.user_id == user_id)
        .order_by(Conversation.updated_at.desc())
    )
    conversations = result.scalars().all()
    return [
        ConversationOut(
            id=str(c.id),
            title=c.title,
            summary=c.summary,
            created_at=c.created_at.isoformat(),
            updated_at=c.updated_at.isoformat(),
        )
        for c in conversations
    ]


@router.post("", response_model=ConversationOut)
async def create_conversation(
    user_id: uuid.UUID = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """Create a new empty conversation."""
    conv = Conversation(user_id=user_id, title="New Chat")
    db.add(conv)
    await db.commit()
    await db.refresh(conv)
    return ConversationOut(
        id=str(conv.id),
        title=conv.title,
        summary=conv.summary,
        created_at=conv.created_at.isoformat(),
        updated_at=conv.updated_at.isoformat(),
    )


@router.delete("/{conversation_id}")
async def delete_conversation(
    conversation_id: uuid.UUID,
    user_id: uuid.UUID = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """Delete a conversation and its messages. RAG embeddings are preserved."""
    result = await db.execute(
        select(Conversation).where(
            Conversation.id == conversation_id,
            Conversation.user_id == user_id,
        )
    )
    conv = result.scalar_one_or_none()
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    # Delete messages belonging to this conversation
    await db.execute(
        delete(Message).where(Message.conversation_id == conversation_id)
    )
    await db.delete(conv)
    await db.commit()
    return {"ok": True}


class SearchResult(BaseModel):
    conversation_id: str
    title: str
    snippet: str
    message_created_at: str


@router.get("/search", response_model=list[SearchResult])
async def search_conversations(
    q: str = Query(..., min_length=1, max_length=200),
    user_id: uuid.UUID = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """Full-text search across user's messages using PostgreSQL tsvector."""
    query = text("""
        SELECT
            m.conversation_id,
            c.title,
            ts_headline('english', m.content, plainto_tsquery('english', :q),
                        'StartSel=**, StopSel=**, MaxWords=30, MinWords=15') AS snippet,
            m.created_at
        FROM messages m
        JOIN conversations c ON c.id = m.conversation_id
        WHERE m.user_id = :user_id
          AND m.conversation_id IS NOT NULL
          AND to_tsvector('english', m.content) @@ plainto_tsquery('english', :q)
        ORDER BY m.created_at DESC
        LIMIT 20
    """)
    result = await db.execute(query, {"q": q, "user_id": user_id})
    rows = result.fetchall()
    return [
        SearchResult(
            conversation_id=str(row.conversation_id),
            title=row.title,
            snippet=row.snippet,
            message_created_at=row.created_at.isoformat(),
        )
        for row in rows
    ]


def _generate_docx(title: str, messages_list: list) -> bytes:
    """Generate a .docx file from conversation messages."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(title, level=1)
    for msg in messages_list:
        ts = msg.created_at.strftime("%Y-%m-%d %H:%M")
        speaker = "User" if msg.role == "user" else "Assistant"
        p = doc.add_paragraph()
        run = p.add_run(f"{speaker} ({ts}): ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p.add_run(msg.content)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_pdf(title: str, messages_list: list) -> bytes:
    """Generate a .pdf file from conversation messages."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
    pdf.add_font("DejaVu", "B", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("DejaVu", size=18)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    for msg in messages_list:
        ts = msg.created_at.strftime("%Y-%m-%d %H:%M")
        speaker = "User" if msg.role == "user" else "Assistant"
        pdf.set_font("DejaVu", style="B", size=10)
        pdf.cell(0, 6, f"{speaker} ({ts}):", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("DejaVu", size=10)
        pdf.multi_cell(0, 5, msg.content)
        pdf.ln(3)
    return bytes(pdf.output())


@router.get("/{conversation_id}/export")
async def export_conversation(
    conversation_id: uuid.UUID,
    format: str = Query("md", pattern="^(md|docx|pdf)$"),
    user_id: uuid.UUID = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """Export a conversation as Markdown, Word, or PDF."""
    result = await db.execute(
        select(Conversation).where(
            Conversation.id == conversation_id,
            Conversation.user_id == user_id,
        )
    )
    conv = result.scalar_one_or_none()
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    msg_result = await db.execute(
        select(Message)
        .where(Message.conversation_id == conversation_id)
        .order_by(Message.created_at)
    )
    messages = msg_result.scalars().all()
    safe_title = conv.title.replace("/", "-").replace("\\", "-")[:80]

    if format == "docx":
        content = _generate_docx(conv.title, messages)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{safe_title}.docx"
    elif format == "pdf":
        content = _generate_pdf(conv.title, messages)
        media_type = "application/pdf"
        filename = f"{safe_title}.pdf"
    else:
        lines = [f"# {conv.title}\n"]
        for msg in messages:
            ts = msg.created_at.strftime("%Y-%m-%d %H:%M")
            speaker = "User" if msg.role == "user" else "Assistant"
            lines.append(f"\n**{speaker}** ({ts}):\n{msg.content}\n")
        content = "\n".join(lines)
        media_type = "text/markdown"
        filename = f"{safe_title}.md"

    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
