"""Generate downloadable document files from text content."""

import csv
import io
import xml.etree.ElementTree as ET

from docx import Document as DocxDocument
from fpdf import FPDF
from openpyxl import Workbook


def generate_document(filename: str, content: str) -> bytes:
    """Generate a document file from text content, dispatching by extension.

    Args:
        filename: Target filename (extension determines format).
        content: The text content to put into the document.

    Returns:
        Raw bytes of the generated document.

    Raises:
        ValueError: If the file extension is not supported.
    """
    ext = _get_extension(filename)
    generators = {
        ".txt": generate_txt,
        ".md": generate_md,
        ".csv": generate_csv,
        ".pdf": generate_pdf,
        ".docx": generate_docx,
        ".xlsx": generate_xlsx,
        ".fdx": generate_fdx,
    }
    gen = generators.get(ext)
    if gen is None:
        supported = ", ".join(sorted(generators.keys()))
        raise ValueError(f"Unsupported file type: {ext}. Supported: {supported}")
    return gen(content)


def generate_txt(content: str) -> bytes:
    """Plain text file."""
    return content.encode("utf-8")


def generate_md(content: str) -> bytes:
    """Markdown file."""
    return content.encode("utf-8")


def generate_csv(content: str) -> bytes:
    """CSV file from tabular content.

    Expects content with rows separated by newlines and columns separated by
    commas, tabs, or pipe characters.
    """
    output = io.StringIO()
    writer = csv.writer(output)
    for line in content.strip().splitlines():
        # Try to detect delimiter
        if "\t" in line:
            cells = line.split("\t")
        elif "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
        else:
            # Use csv reader to handle quoted commas
            cells = next(csv.reader(io.StringIO(line)))
        writer.writerow(cells)
    return output.getvalue().encode("utf-8")


def generate_pdf(content: str) -> bytes:
    """PDF file with text content."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    # Helvetica only supports Latin-1; replace common unicode chars with ASCII
    safe_content = _sanitize_for_latin1(content)
    pdf.multi_cell(0, 7, safe_content)
    return bytes(pdf.output())


# Unicode → ASCII replacements for Latin-1 safe PDF output
_UNICODE_REPLACEMENTS = {
    "\u2013": "-",   # en-dash
    "\u2014": "--",  # em-dash
    "\u2018": "'",   # left single quote
    "\u2019": "'",   # right single quote
    "\u201c": '"',   # left double quote
    "\u201d": '"',   # right double quote
    "\u2026": "...", # ellipsis
    "\u2022": "*",   # bullet
    "\u00a0": " ",   # non-breaking space
    "\u200b": "",    # zero-width space
    "\u2011": "-",   # non-breaking hyphen
    "\u2010": "-",   # hyphen
    "\ufeff": "",    # BOM
}


def _sanitize_for_latin1(text: str) -> str:
    """Replace unicode characters that Helvetica cannot render."""
    for char, replacement in _UNICODE_REPLACEMENTS.items():
        text = text.replace(char, replacement)
    # Drop any remaining non-Latin-1 characters
    return text.encode("latin-1", errors="replace").decode("latin-1")


def generate_docx(content: str) -> bytes:
    """Word document with paragraphs."""
    doc = DocxDocument()
    for paragraph in content.split("\n"):
        doc.add_paragraph(paragraph)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_xlsx(content: str) -> bytes:
    """Excel spreadsheet from tabular content.

    Expects rows separated by newlines and columns separated by commas, tabs,
    or pipe characters.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    for line in content.strip().splitlines():
        if "\t" in line:
            cells = line.split("\t")
        elif "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
        else:
            cells = next(csv.reader(io.StringIO(line)))
        ws.append(cells)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def generate_fdx(content: str) -> bytes:
    """Final Draft XML (.fdx) from screenplay-formatted text.

    Expects standard screenplay formatting:
    - Lines in ALL CAPS starting with INT./EXT. → Scene Heading
    - Lines in ALL CAPS (short, not scene headings) → Character
    - Lines in (parentheses) → Parenthetical
    - Lines starting with > or ending with : in ALL CAPS → Transition
    - Indented lines after a character → Dialogue
    - Everything else → Action
    """
    root = ET.Element("FinalDraft", DocumentType="Script", Template="No", Version="5")
    content_el = ET.SubElement(root, "Content")

    lines = content.strip().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Scene heading: INT. / EXT. / I/E.
        if stripped.upper().startswith(("INT.", "EXT.", "INT ", "EXT ", "I/E.")):
            _add_paragraph(content_el, "Scene Heading", stripped.upper())
            i += 1
        # Transition: ends with "TO:" or starts with ">"
        elif stripped.startswith(">") or (
            stripped.upper() == stripped and stripped.endswith(":")
        ):
            text = stripped.lstrip(">").strip()
            _add_paragraph(content_el, "Transition", text.upper())
            i += 1
        # Character name: all caps, relatively short, next line is dialogue
        elif (
            stripped.upper() == stripped
            and len(stripped) < 50
            and stripped.isalpha() or stripped.replace(" ", "").replace("(", "").replace(")", "").isalpha()
        ) and stripped.upper() == stripped and len(stripped) > 0:
            _add_paragraph(content_el, "Character", stripped)
            i += 1
            # Consume following parenthetical and dialogue lines
            while i < len(lines):
                next_line = lines[i].rstrip()
                next_stripped = next_line.strip()
                if not next_stripped:
                    break
                if next_stripped.startswith("(") and next_stripped.endswith(")"):
                    _add_paragraph(content_el, "Parenthetical", next_stripped)
                elif not (next_stripped.upper() == next_stripped and len(next_stripped) < 50):
                    _add_paragraph(content_el, "Dialogue", next_stripped)
                else:
                    break
                i += 1
        else:
            # Action
            _add_paragraph(content_el, "Action", stripped)
            i += 1

    tree = ET.ElementTree(root)
    buf = io.BytesIO()
    tree.write(buf, encoding="utf-8", xml_declaration=True)
    return buf.getvalue()


def _add_paragraph(parent: ET.Element, ptype: str, text: str) -> None:
    """Add a Paragraph element with a Text child to the FDX XML."""
    para = ET.SubElement(parent, "Paragraph", Type=ptype)
    text_el = ET.SubElement(para, "Text")
    text_el.text = text


def _get_extension(filename: str) -> str:
    """Return the lowercase file extension including the dot."""
    dot = filename.rfind(".")
    if dot == -1:
        return ""
    return filename[dot:].lower()
