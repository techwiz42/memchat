"""Format-preserving document editor.

Given original file bytes and new text content from the LLM, applies text-level
changes to the original document while preserving its structure and formatting.
"""

import io
import logging
import os
import re
import subprocess
import tempfile
import xml.etree.ElementTree as ET
from copy import deepcopy
from difflib import SequenceMatcher

import pdfplumber

# XML-illegal control characters: 0x00-0x08, 0x0B, 0x0C, 0x0E-0x1F, 0x7F
# Tab (0x09), newline (0x0A), carriage return (0x0D) are allowed.
_XML_ILLEGAL_RE = re.compile(
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]"
)


def _sanitize_for_xml(text: str) -> str:
    """Strip characters that are illegal in XML from text."""
    return _XML_ILLEGAL_RE.sub("", text)
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from pdf2docx import Converter

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Plain-text / Markdown section editing and find-replace
# ---------------------------------------------------------------------------

def edit_text_section(
    original_bytes: bytes,
    sections_json: list[dict],
    section_index: int,
    new_content: str,
) -> bytes | None:
    """Edit a section in a plain text or markdown file.

    Decodes to text, splits on paragraph boundaries (\\n\\n), splices in
    the new content at the section's paragraph range, and re-encodes.
    """
    if section_index < 0 or section_index >= len(sections_json):
        logger.warning("edit_text_section: invalid section_index %d", section_index)
        return None

    new_content = _sanitize_for_xml(new_content)
    text = original_bytes.decode("utf-8", errors="replace")
    paragraphs = text.split("\n\n")

    section = sections_json[section_index]
    para_start = section["xml_para_start"]
    para_end = section["xml_para_end"]

    # Clamp to actual paragraph count
    para_start = max(0, min(para_start, len(paragraphs)))
    para_end = max(para_start, min(para_end, len(paragraphs)))

    new_paragraphs = new_content.split("\n\n")
    edited = paragraphs[:para_start] + new_paragraphs + paragraphs[para_end:]
    return "\n\n".join(edited).encode("utf-8")


def find_replace_text(
    original_bytes: bytes,
    find: str,
    replace: str,
    case_sensitive: bool = False,
) -> tuple[bytes, int] | None:
    """Global find-and-replace in a plain text file.

    Returns (edited_bytes, replacement_count), or None if no matches.
    """
    text = original_bytes.decode("utf-8", errors="replace")
    flags = 0 if case_sensitive else re.IGNORECASE
    count = len(re.findall(re.escape(find), text, flags=flags))
    if count == 0:
        return None
    edited = re.sub(re.escape(find), replace, text, flags=flags)
    return edited.encode("utf-8"), count


def find_replace_docx(
    original_bytes: bytes,
    find: str,
    replace: str,
    case_sensitive: bool = False,
) -> tuple[bytes, int] | None:
    """Global find-and-replace directly on DOCX paragraph runs.

    Much faster than find_replace_rich because it operates on the DOCX
    object model directly without full-text extraction and SequenceMatcher.
    """
    try:
        doc = DocxDocument(io.BytesIO(original_bytes))
    except Exception:
        logger.warning("find_replace_docx: failed to parse DOCX")
        return None

    replace = _sanitize_for_xml(replace)
    flags = 0 if case_sensitive else re.IGNORECASE
    pattern = re.compile(re.escape(find), flags)
    count = 0

    def _replace_in_paragraph(para):
        nonlocal count
        text = para.text
        if not text:
            return
        matches = len(pattern.findall(text))
        if matches == 0:
            return
        count += matches
        new_text = pattern.sub(replace, text)
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.text = new_text

    for para in doc.paragraphs:
        _replace_in_paragraph(para)

    # Also handle table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)

    if count == 0:
        return None

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), count


# ---------------------------------------------------------------------------
# Rich format (DOCX / PDF) section editing and find-replace
# ---------------------------------------------------------------------------

def edit_docx_section(
    original_bytes: bytes,
    sections_json: list[dict],
    section_index: int,
    new_content: str,
) -> bytes | None:
    """Surgically replace paragraphs in one section of a DOCX file.

    Maps the section's paragraph range (indices into extracted non-empty
    paragraphs) back to actual python-docx paragraph objects and replaces
    their text directly.  Handles insert (new paragraphs) and delete
    (removed paragraphs) via underlying XML manipulation.

    Much faster than edit_rich_section because it avoids full-document
    text extraction and SequenceMatcher diffing on every edit.
    """
    new_content = _sanitize_for_xml(new_content)
    if section_index < 0 or section_index >= len(sections_json):
        logger.warning("edit_docx_section: invalid section_index %d", section_index)
        return None

    try:
        doc = DocxDocument(io.BytesIO(original_bytes))
    except Exception:
        logger.warning("edit_docx_section: failed to parse DOCX")
        return None

    # Build mapping: extracted-text paragraph index → docx paragraph object.
    # _extract_docx filters to non-empty paragraphs, so this mapping aligns
    # with the indices produced by split_large_text.
    non_empty: list[int] = []  # docx paragraph indices with non-empty text
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            non_empty.append(i)

    section = sections_json[section_index]
    para_start = section["xml_para_start"]
    para_end = section["xml_para_end"]

    # Clamp
    para_start = max(0, min(para_start, len(non_empty)))
    para_end = max(para_start, min(para_end, len(non_empty)))

    target_docx_indices = non_empty[para_start:para_end]

    # Parse new content into paragraphs (split on double-newline, matching
    # the same convention used by split_large_text)
    new_paras = [p.strip() for p in new_content.split("\n\n") if p.strip()]

    # Apply changes: replace existing, insert new, delete removed
    n_orig = len(target_docx_indices)
    n_new = len(new_paras)
    common = min(n_orig, n_new)

    # 1) Replace text in existing paragraphs (preserves styles)
    for k in range(common):
        para = doc.paragraphs[target_docx_indices[k]]
        if para.runs:
            para.runs[0].text = new_paras[k]
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.text = new_paras[k]

    # 2) Extra new paragraphs → insert after the last replaced paragraph
    if n_new > common:
        if common > 0:
            anchor_el = doc.paragraphs[target_docx_indices[common - 1]]._element
        elif para_start > 0:
            # Insert before the section's start (use preceding paragraph)
            anchor_el = doc.paragraphs[non_empty[para_start - 1]]._element
        else:
            anchor_el = doc.paragraphs[0]._element

        for k in range(common, n_new):
            # Clone the anchor paragraph to inherit its style, then set text
            new_p = deepcopy(anchor_el)
            # Clear all run text then set the first run
            for r in new_p.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    t.text = ""
            # Set text in first run, or create one
            runs = new_p.findall(qn("w:r"))
            if runs:
                t_els = runs[0].findall(qn("w:t"))
                if t_els:
                    t_els[0].text = new_paras[k]
                else:
                    t_el = ET.SubElement(runs[0], qn("w:t"))
                    t_el.text = new_paras[k]
            else:
                r_el = ET.SubElement(new_p, qn("w:r"))
                t_el = ET.SubElement(r_el, qn("w:t"))
                t_el.text = new_paras[k]

            anchor_el.addnext(new_p)
            anchor_el = new_p  # chain insertions in order

    # 3) Extra original paragraphs → clear them (effectively deleting content)
    for k in range(common, n_orig):
        para = doc.paragraphs[target_docx_indices[k]]
        # Clear text rather than removing the XML element, which could break
        # document structure (e.g., section breaks attached to paragraphs)
        if para.runs:
            para.runs[0].text = ""
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.text = ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def edit_rich_section(
    original_bytes: bytes,
    original_filename: str,
    sections_json: list[dict],
    section_index: int,
    new_content: str,
    extract_fn,
) -> bytes | None:
    """Edit a section in a PDF by extracting text, splicing paragraphs,
    then reconstructing via edit_preserving_format().

    For DOCX files, use edit_docx_section() instead — it's much faster.
    extract_fn should be parser.extract_text_sync.
    """
    new_content = _sanitize_for_xml(new_content)
    if section_index < 0 or section_index >= len(sections_json):
        logger.warning("edit_rich_section: invalid section_index %d", section_index)
        return None

    full_text = extract_fn(original_filename, original_bytes)
    paragraphs = full_text.split("\n\n")

    section = sections_json[section_index]
    para_start = section["xml_para_start"]
    para_end = section["xml_para_end"]

    para_start = max(0, min(para_start, len(paragraphs)))
    para_end = max(para_start, min(para_end, len(paragraphs)))

    new_paragraphs = new_content.split("\n\n")
    edited_paragraphs = paragraphs[:para_start] + new_paragraphs + paragraphs[para_end:]
    edited_text = "\n\n".join(edited_paragraphs)

    return edit_preserving_format(original_bytes, original_filename, edited_text)


def find_replace_rich(
    original_bytes: bytes,
    original_filename: str,
    find: str,
    replace: str,
    case_sensitive: bool = False,
    extract_fn=None,
) -> tuple[bytes, int] | None:
    """Global find-and-replace for DOCX/PDF.

    Extracts text, performs replacement, and reconstructs the file
    via edit_preserving_format().  Returns (edited_bytes, count) or None.
    """
    replace = _sanitize_for_xml(replace)
    full_text = extract_fn(original_filename, original_bytes)
    flags = 0 if case_sensitive else re.IGNORECASE
    count = len(re.findall(re.escape(find), full_text, flags=flags))
    if count == 0:
        return None
    edited_text = re.sub(re.escape(find), replace, full_text, flags=flags)
    edited_bytes = edit_preserving_format(original_bytes, original_filename, edited_text)
    if edited_bytes is None:
        return None
    return edited_bytes, count


def edit_preserving_format(
    original_bytes: bytes,
    original_filename: str,
    new_text: str,
) -> bytes | None:
    """Apply text edits to the original file, preserving formatting.

    Returns edited file bytes, or None if format-preserving edit is not
    supported for this file type (caller should fall back to generate_document).
    """
    ext = _get_extension(original_filename)
    if ext == ".fdx":
        return _edit_fdx(original_bytes, new_text)
    if ext == ".docx":
        return _edit_docx(original_bytes, new_text)
    if ext == ".pdf":
        return _edit_pdf(original_bytes, new_text)
    return None


# ---------------------------------------------------------------------------
# PDF editing via PDF → DOCX → edit → PDF pipeline
# ---------------------------------------------------------------------------

def _edit_pdf(original_bytes: bytes, new_text: str) -> bytes | None:
    """Edit a PDF by converting to DOCX, applying edits, converting back.

    The LLM sees pdfplumber-extracted text, but the intermediate DOCX comes from
    pdf2docx which produces different paragraph structure.  We bridge the gap by:
    1. Re-extracting text with pdfplumber (matching what the LLM saw)
    2. Diffing that against the LLM's new text to find specific changes
    3. Applying those changes as find-and-replace on the DOCX paragraphs
    """
    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = os.path.join(tmp, "original.pdf")
        docx_path = os.path.join(tmp, "converted.docx")
        out_pdf_path = os.path.join(tmp, "edited.pdf")

        # 1. Write original PDF
        with open(pdf_path, "wb") as f:
            f.write(original_bytes)

        # 2. Re-extract text with pdfplumber (same method used at upload time)
        original_extracted = _extract_pdf_text(original_bytes)

        # 3. Compute specific text changes between extraction and LLM's output
        replacements = _compute_text_replacements(original_extracted, new_text)
        if not replacements:
            logger.info("No text changes detected for PDF edit")
            return None

        logger.info("PDF edit: found %d text replacement(s)", len(replacements))

        # 4. PDF → DOCX
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
        except Exception as e:
            logger.warning("pdf2docx conversion failed: %s", e)
            return None

        # 5. Apply replacements to the DOCX
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()

        edited_docx = _apply_replacements_to_docx(docx_bytes, replacements)
        if edited_docx is None:
            return None

        edited_docx_path = os.path.join(tmp, "edited.docx")
        with open(edited_docx_path, "wb") as f:
            f.write(edited_docx)

        # 6. DOCX → PDF via LibreOffice
        try:
            subprocess.run(
                [
                    "libreoffice", "--headless", "--norestore",
                    "--convert-to", "pdf",
                    "--outdir", tmp,
                    edited_docx_path,
                ],
                check=True,
                capture_output=True,
                timeout=60,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
            logger.warning("LibreOffice PDF conversion failed: %s", e)
            return None

        if not os.path.exists(out_pdf_path):
            # LibreOffice names output after the input file
            alt_path = os.path.join(tmp, "edited.pdf")
            if not os.path.exists(alt_path):
                logger.warning("LibreOffice did not produce output PDF")
                return None
            out_pdf_path = alt_path

        with open(out_pdf_path, "rb") as f:
            return f.read()


def _extract_pdf_text(pdf_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber (same method as parser.py)."""
    pages = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def _compute_text_replacements(
    original: str, new: str,
) -> list[tuple[str, str]]:
    """Diff original vs new text at the word level, returning (old, new) phrase pairs."""
    orig_words = original.split()
    new_words = new.split()

    matcher = SequenceMatcher(
        None, [w.lower() for w in orig_words], [w.lower() for w in new_words],
    )

    replacements: list[tuple[str, str]] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        old_phrase = " ".join(orig_words[i1:i2])
        new_phrase = " ".join(new_words[j1:j2])
        if old_phrase != new_phrase:
            replacements.append((old_phrase, new_phrase))
            logger.info("PDF edit change: %r -> %r", old_phrase[:80], new_phrase[:80])

    return replacements


def _apply_replacements_to_docx(
    docx_bytes: bytes, replacements: list[tuple[str, str]],
) -> bytes | None:
    """Apply find-and-replace text changes to a DOCX, preserving formatting."""
    try:
        doc = DocxDocument(io.BytesIO(docx_bytes))
    except Exception:
        logger.warning("Failed to parse DOCX for replacement")
        return None

    _WS = re.compile(r"\s+")
    applied_count = 0

    def _apply_to_paragraph(para):
        """Apply replacements to a single paragraph, return True if changed."""
        text = para.text
        if not text.strip():
            return False

        new_text_val = text
        for old_phrase, new_phrase in replacements:
            if not old_phrase:
                continue
            # Normalize whitespace for matching (pdf2docx may use different spacing)
            pattern = _WS.sub(r"\\s+", re.escape(old_phrase))
            match = re.search(pattern, new_text_val)
            if match:
                logger.info("MATCH: %r -> %r", old_phrase[:60], new_phrase[:60])
                new_text_val = new_text_val[:match.start()] + new_phrase + new_text_val[match.end():]

        if new_text_val != text:
            if para.runs:
                para.runs[0].text = new_text_val
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = new_text_val
            return True
        return False

    # Apply to body paragraphs
    for para in doc.paragraphs:
        if _apply_to_paragraph(para):
            applied_count += 1

    # Apply to table cells (pdf2docx often puts structured data in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _apply_to_paragraph(para):
                        applied_count += 1

    logger.info("Applied replacements to %d paragraph(s)", applied_count)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# FDX global find-and-replace
# ---------------------------------------------------------------------------

def find_replace_fdx(
    original_bytes: bytes,
    find: str,
    replace: str,
    case_sensitive: bool = False,
) -> tuple[bytes, int] | None:
    """Global find-and-replace across all Text elements in an FDX file.

    Returns (edited_bytes, replacement_count), or None on parse failure.
    """
    try:
        root = ET.fromstring(original_bytes)
    except ET.ParseError:
        logger.warning("find_replace_fdx: failed to parse FDX XML")
        return None

    count = 0
    for text_el in root.iter("Text"):
        if text_el.text:
            if case_sensitive:
                new_text = text_el.text.replace(find, replace)
            else:
                # Case-insensitive replace preserving surrounding case
                import re as _re
                new_text = _re.sub(_re.escape(find), replace, text_el.text, flags=_re.IGNORECASE)
            if new_text != text_el.text:
                occurrences = len(text_el.text.split(find if case_sensitive else find.lower())) - 1
                if not case_sensitive:
                    occurrences = len(_re.findall(_re.escape(find), text_el.text, flags=_re.IGNORECASE))
                count += occurrences
                text_el.text = new_text

    if count == 0:
        return None

    buf = io.BytesIO()
    tree = ET.ElementTree(root)
    tree.write(buf, encoding="utf-8", xml_declaration=True)
    return buf.getvalue(), count


# ---------------------------------------------------------------------------
# FDX section editing (surgical single-scene replacement)
# ---------------------------------------------------------------------------

def edit_fdx_section(
    original_bytes: bytes,
    sections_json: list[dict],
    section_index: int,
    new_content: str,
) -> bytes | None:
    """Surgically replace paragraphs in one scene of an FDX file.

    Uses the xml_para_start:xml_para_end range from sections_json to identify
    the target paragraphs, then applies the SequenceMatcher approach only to
    that range while leaving all other scenes untouched.

    Returns the edited FDX bytes, or None on failure.
    """
    new_content = _sanitize_for_xml(new_content)
    if section_index < 0 or section_index >= len(sections_json):
        logger.warning("edit_fdx_section: invalid section_index %d", section_index)
        return None

    try:
        root = ET.fromstring(original_bytes)
    except ET.ParseError:
        logger.warning("edit_fdx_section: failed to parse FDX XML")
        return None

    content_el = root.find("Content")
    if content_el is None:
        return None

    paragraphs = list(content_el.findall("Paragraph"))
    section = sections_json[section_index]
    para_start = section["xml_para_start"]
    para_end = section["xml_para_end"]

    # Clamp range to actual paragraph count
    para_start = max(0, min(para_start, len(paragraphs)))
    para_end = max(para_start, min(para_end, len(paragraphs)))

    # Extract original paragraph texts in the target range
    orig_range = paragraphs[para_start:para_end]
    orig_texts_lower = []
    for p in orig_range:
        texts = []
        for t in p.iter("Text"):
            if t.text:
                texts.append(t.text)
        orig_texts_lower.append("".join(texts).strip().lower())

    # Parse new content into lines
    new_lines = [line.strip() for line in new_content.splitlines() if line.strip()]
    new_lower = [line.lower() for line in new_lines]

    # Use SequenceMatcher to align original paragraphs with new text
    matcher = SequenceMatcher(None, orig_texts_lower, new_lower)
    new_paras: list[ET.Element] = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for idx in range(i1, i2):
                new_paras.append(orig_range[idx])
        elif tag == "replace":
            orig_slice = list(range(i1, i2))
            new_slice = list(range(j1, j2))
            for k, j_idx in enumerate(new_slice):
                if k < len(orig_slice):
                    para_el = orig_range[orig_slice[k]]
                    _set_paragraph_text(para_el, new_lines[j_idx])
                    new_paras.append(para_el)
                else:
                    new_paras.append(_make_paragraph("Action", new_lines[j_idx]))
        elif tag == "insert":
            for j_idx in range(j1, j2):
                new_paras.append(_make_paragraph("Action", new_lines[j_idx]))
        elif tag == "delete":
            pass  # lines removed by the edit

    # Rebuild Content: paragraphs before section + edited + paragraphs after
    new_content_el = ET.Element("Content")
    for p in paragraphs[:para_start]:
        new_content_el.append(p)
    for p in new_paras:
        new_content_el.append(p)
    for p in paragraphs[para_end:]:
        new_content_el.append(p)

    root.remove(content_el)
    root.append(new_content_el)

    buf = io.BytesIO()
    tree = ET.ElementTree(root)
    tree.write(buf, encoding="utf-8", xml_declaration=True)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# FDX editing (Final Draft screenplay XML) — full document
# ---------------------------------------------------------------------------

def _edit_fdx(original_bytes: bytes, new_text: str) -> bytes | None:
    """Edit an FDX screenplay XML, replacing text while preserving structure."""
    new_text = _sanitize_for_xml(new_text)
    try:
        root = ET.fromstring(original_bytes)
    except ET.ParseError:
        logger.warning("Failed to parse original FDX XML")
        return None

    # Collect original paragraph elements and their plain text
    content_el = root.find("Content")
    if content_el is None:
        return None

    orig_paragraphs: list[tuple[ET.Element, str]] = []
    for para in content_el.findall("Paragraph"):
        texts = []
        for text_el in para.iter("Text"):
            if text_el.text:
                texts.append(text_el.text)
        line = "".join(texts).strip()
        if line:
            orig_paragraphs.append((para, line))

    # Parse new text into lines (strip screenplay indentation from extraction)
    new_lines = [
        line.strip() for line in new_text.splitlines() if line.strip()
    ]

    # Build case-insensitive versions for matching
    orig_texts = [t.lower() for _, t in orig_paragraphs]
    new_lower = [t.lower() for t in new_lines]

    # Use SequenceMatcher to align original paragraphs with new text
    matcher = SequenceMatcher(None, orig_texts, new_lower)
    opcodes = matcher.get_opcodes()

    # Build a new Content element preserving structure
    new_content = ET.Element("Content")

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            # Keep original elements unchanged (they match)
            for idx in range(i1, i2):
                new_content.append(orig_paragraphs[idx][0])
        elif tag == "replace":
            # Map new text onto original elements where possible
            orig_slice = list(range(i1, i2))
            new_slice = list(range(j1, j2))
            for k, j_idx in enumerate(new_slice):
                if k < len(orig_slice):
                    # Update text in existing element (preserves Type attribute)
                    para_el = orig_paragraphs[orig_slice[k]][0]
                    _set_paragraph_text(para_el, new_lines[j_idx])
                    new_content.append(para_el)
                else:
                    # Extra new lines — create Action paragraphs
                    new_content.append(_make_paragraph("Action", new_lines[j_idx]))
            # Original paragraphs beyond the new slice are dropped (deleted)
        elif tag == "insert":
            # New lines with no original counterpart
            for j_idx in range(j1, j2):
                new_content.append(_make_paragraph("Action", new_lines[j_idx]))
        elif tag == "delete":
            # Original lines removed in the edit — skip them
            pass

    # Replace Content element in the tree
    root.remove(content_el)
    root.append(new_content)

    buf = io.BytesIO()
    tree = ET.ElementTree(root)
    tree.write(buf, encoding="utf-8", xml_declaration=True)
    return buf.getvalue()


def _set_paragraph_text(para: ET.Element, new_text: str) -> None:
    """Replace all Text element content in a Paragraph with new text."""
    text_els = list(para.iter("Text"))
    if text_els:
        # Put all text in the first Text element, clear the rest
        text_els[0].text = new_text
        for extra in text_els[1:]:
            extra.text = ""
    else:
        text_el = ET.SubElement(para, "Text")
        text_el.text = new_text


def _make_paragraph(ptype: str, text: str) -> ET.Element:
    """Create a new FDX Paragraph element."""
    para = ET.Element("Paragraph", Type=ptype)
    text_el = ET.SubElement(para, "Text")
    text_el.text = text
    return para


# ---------------------------------------------------------------------------
# DOCX editing
# ---------------------------------------------------------------------------

def _edit_docx(original_bytes: bytes, new_text: str) -> bytes | None:
    """Edit a DOCX file, replacing paragraph text while preserving styles."""
    new_text = _sanitize_for_xml(new_text)
    try:
        doc = DocxDocument(io.BytesIO(original_bytes))
    except Exception:
        logger.warning("Failed to parse original DOCX")
        return None

    # Extract original paragraph texts (non-empty only)
    orig_paras_with_idx: list[tuple[int, str]] = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            orig_paras_with_idx.append((i, para.text.strip()))

    new_lines = [line.strip() for line in new_text.splitlines() if line.strip()]

    orig_texts = [t.lower() for _, t in orig_paras_with_idx]
    new_lower = [t.lower() for t in new_lines]

    matcher = SequenceMatcher(None, orig_texts, new_lower)
    opcodes = matcher.get_opcodes()

    # Track which original paragraphs to update
    updates: dict[int, str] = {}  # doc paragraph index -> new text

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            pass
        elif tag == "replace":
            orig_slice = list(range(i1, i2))
            new_slice = list(range(j1, j2))
            for k, j_idx in enumerate(new_slice):
                if k < len(orig_slice):
                    doc_idx = orig_paras_with_idx[orig_slice[k]][0]
                    updates[doc_idx] = new_lines[j_idx]

    # Apply updates: replace text while preserving the first run's formatting
    for doc_idx, new_text_val in updates.items():
        para = doc.paragraphs[doc_idx]
        if para.runs:
            para.runs[0].text = new_text_val
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.text = new_text_val

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_extension(filename: str) -> str:
    dot = filename.rfind(".")
    if dot == -1:
        return ""
    return filename[dot:].lower()
